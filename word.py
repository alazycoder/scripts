# coding=utf-8
import docx
import sys
import urllib
import traceback
from selenium import webdriver
from bs4 import BeautifulSoup


def get_soup(url):
    page_source = get_html_by_phantomjs(url)
    soup = BeautifulSoup(page_source, "html.parser")
    return soup


def get_html_by_phantomjs(url):
    try:
        driver = webdriver.PhantomJS(executable_path="/Users/gaohongjie/phantomjs/bin/phantomjs")
        driver.get(url)
        return driver.page_source
    except:
        return "error"


if __name__ == "__main__":
    reload(sys)
    sys.setdefaultencoding('utf-8')
    out = docx.Document("/Users/gaohongjie/Desktop/out.docx")
    f = docx.Document("/Users/gaohongjie/Desktop/in.docx")
    for i in range(63, len(f.paragraphs)):
        try:
            s = str(f.paragraphs[i].text)
            print i
            print s
            url = "https://translate.google.cn/#en/zh-CN/" + urllib.quote(s)
            soup = get_soup(url)
            if soup.text == 'error':
                print "第 %d 段翻译出错,段落内容为 %s" % (i, s)
                continue
            span = soup.find('span', {'id': 'result_box'})
            for item in span.children:
                if item.name == "span":
                    print item.text
                    out.add_paragraph(item.text)
                    out.save('/Users/gaohongjie/Desktop/out.docx')
            print "第 %d 段翻译成功" % i
        except:
            traceback.print_exc()
            print "第 %d 段翻译出错" % i
